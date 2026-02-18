#!/usr/bin/env python3
"""
word_utils.py — cross-platform Word table population utility for ClipCommand.

Drop this file into your transforms/ folder alongside the word_from_* scripts.
Import it in any transform:
    from word_utils import update_table, WordUtilsError

Public API
----------
update_table(bookmark, data, heading_rows=1, doc_path=None)
    Write *data* (list of lists) into the first table inside *bookmark*.

    bookmark      : str   — Word bookmark name (no spaces)
    data          : list  — list of rows, each row a list of cell values
    heading_rows  : int   — how many heading rows already exist in the table
                            (these are skipped; data is written below them)
    doc_path      : str   — path to .docx (required for python-docx fallback only;
                            ignored on Windows COM and macOS AppleScript paths,
                            which both operate on the currently active document)

Platform dispatch
-----------------
  Windows  → win32com COM automation  (live active document, no save needed)
  macOS    → AppleScript via osascript (live active document, no save needed)
  Linux /
  fallback → python-docx              (file on disk; doc_path required;
                                       Word must NOT have the file open)
"""

import subprocess
import sys
import os
from pathlib import Path


class WordUtilsError(Exception):
    """Raised for all word_utils failures."""


# ─── Public entry point ───────────────────────────────────────────────────────

def update_table(bookmark: str, data: list, heading_rows: int = 1,
                 doc_path: str = None) -> str:
    """
    Write data into the bookmarked Word table.
    Returns a short status string suitable for use as transform() return value.
    """
    if not data:
        raise WordUtilsError("No data rows provided.")

    if sys.platform == "win32":
        return _update_via_com(bookmark, data, heading_rows)
    elif sys.platform == "darwin":
        return _update_via_applescript(bookmark, data, heading_rows)
    else:
        if not doc_path:
            raise WordUtilsError(
                "doc_path is required on Linux / non-Windows platforms."
            )
        return _update_via_docx(bookmark, data, heading_rows, doc_path)


# ─── Windows: win32com ────────────────────────────────────────────────────────

def _get_com_word():
    """Return a live Word.Application COM object, clearing the gen_py cache if needed."""
    import shutil
    from win32com import client
    try:
        return client.gencache.EnsureDispatch("Word.Application")
    except AttributeError:
        import re
        for mod in list(sys.modules):
            if re.match(r"win32com\.gen_py\..+", mod):
                del sys.modules[mod]
        shutil.rmtree(
            os.path.join(os.environ.get("LOCALAPPDATA"), "Temp", "gen_py"),
            ignore_errors=True,
        )
        return client.gencache.EnsureDispatch("Word.Application")


def _update_via_com(bookmark: str, data: list, heading_rows: int) -> str:
    try:
        app = _get_com_word()
        app.Visible       = True
        app.DisplayAlerts = False
        doc               = app.ActiveDocument

        if doc is None:
            raise WordUtilsError("No Word document is currently open.")

        word_range  = doc.Bookmarks(bookmark).Range
        table       = word_range.Tables(1)
        rows_needed = len(data) + heading_rows

        if table.Rows.Count < rows_needed:
            table.Select()
            app.Selection.InsertRowsBelow(NumRows=rows_needed - table.Rows.Count)

        for row_idx, entry in enumerate(data, start=heading_rows + 1):
            for col_idx, value in enumerate(entry, start=1):
                table.Cell(row_idx, col_idx).Range.Text = str(value)

        return (
            f"OK — {len(data)} row(s) written to bookmark '{bookmark}' "
            f"in '{doc.Name}' (Windows COM)"
        )
    except WordUtilsError:
        raise
    except Exception as exc:
        raise WordUtilsError(f"COM error: {exc}") from exc


# ─── macOS: AppleScript ───────────────────────────────────────────────────────
#
# Strategy:
#   1. Find the first table inside the bookmark range.
#   2. Count existing rows; add rows if needed using "insert rows below".
#   3. Set cell content row by row using "set content of text object of cell".
#
# The data is passed as a JSON-encoded string argument so we avoid any
# AppleScript string-escaping nightmares with arbitrary cell values.
# The AppleScript reads the JSON file, iterates, and writes cells.
# We write the data to a temp file and pass the path to osascript.

def _update_via_applescript(bookmark: str, data: list, heading_rows: int) -> str:
    import json
    import tempfile

    # Serialise data to a temp JSON file — sidesteps all quoting issues
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".json", delete=False, encoding="utf-8"
    ) as tf:
        json.dump(data, tf, ensure_ascii=False)
        tmp_path = tf.name

    # Build the rows as AppleScript list-of-lists string
    # We pass file path and params as positional argv to osascript
    script = f"""
use AppleScript version "2.4"
use scripting additions

-- Read data from temp JSON via shell (avoids escaping nightmares)
set jsonPath to "{tmp_path}"
set jsonText to do shell script "cat " & quoted form of jsonPath

-- Parse JSON into AS list using Python (available on all macOS)
set pyCmd to "import sys,json; d=json.load(open(sys.argv[1])); " & ¬
    "print('\\n'.join(['\\t'.join(str(c) for c in r) for r in d]))"
set tabData to do shell script "python3 -c " & quoted form of pyCmd & " " & quoted form of jsonPath

-- Split into rows then cells
set rowList to paragraphs of tabData
set headingRows to {heading_rows}
set bookmarkName to "{bookmark}"

tell application "Microsoft Word"
    if not (exists active document) then
        error "No Word document is currently open."
    end if

    tell active document
        -- Locate the bookmark
        if not (exists bookmark bookmarkName) then
            error "Bookmark '" & bookmarkName & "' not found in active document."
        end if

        set bkRange to text object of bookmark bookmarkName
        set tbl to table 1 of bkRange

        set existingRows to count rows of tbl
        set neededRows to (count of rowList) + headingRows

        -- Add rows if the table is too short
        if existingRows < neededRows then
            set extraRows to neededRows - existingRows
            insert rows below (last row of tbl) count extraRows
        end if

        -- Write cell values
        set rowIdx to headingRows + 1
        repeat with aRow in rowList
            set cellValues to every text item of aRow
            set colIdx to 1
            repeat with aCell in cellValues
                try
                    set content of text object of cell rowIdx colIdx of tbl to (aCell as text)
                end try
                set colIdx to colIdx + 1
            end repeat
            set rowIdx to rowIdx + 1
        end repeat

    end tell
end tell

return "OK"
"""

    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            err = result.stderr.strip()
            raise WordUtilsError(f"AppleScript error: {err}")

        return (
            f"OK — {len(data)} row(s) written to bookmark '{bookmark}' "
            f"in active document (macOS AppleScript)"
        )
    except WordUtilsError:
        raise
    except subprocess.TimeoutExpired:
        raise WordUtilsError("AppleScript timed out after 60s.")
    except Exception as exc:
        raise WordUtilsError(f"osascript error: {exc}") from exc
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ─── Linux / fallback: python-docx ───────────────────────────────────────────
#
# Operates on the .docx file on disk.  The file must NOT be open in Word.
# Bookmarks are found by scanning the XML for w:bookmarkStart elements,
# then walking up to find the enclosing table.

def _update_via_docx(bookmark: str, data: list, heading_rows: int,
                     doc_path: str) -> str:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise WordUtilsError(
            "python-docx is not installed. pip install python-docx"
        )

    path = Path(doc_path).expanduser().resolve()
    if not path.exists():
        raise WordUtilsError(f"Document not found: {path}")

    doc = Document(str(path))

    # ── Find the bookmark and its enclosing table ─────────────────────────────
    table = _docx_find_table_for_bookmark(doc, bookmark)
    if table is None:
        raise WordUtilsError(
            f"Bookmark '{bookmark}' not found, or not inside a table."
        )

    # ── Add rows if needed ────────────────────────────────────────────────────
    rows_needed = len(data) + heading_rows
    while len(table.rows) < rows_needed:
        _docx_add_row(table)

    # ── Write cell values ─────────────────────────────────────────────────────
    for row_idx, entry in enumerate(data, start=heading_rows):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(entry):
            if col_idx < len(row.cells):
                # Clear existing paragraphs and set text cleanly
                cell = row.cells[col_idx]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                if cell.paragraphs:
                    cell.paragraphs[0].runs[0].text if cell.paragraphs[0].runs else None
                    # Use the cell's text property directly
                cell.text = str(value)

    doc.save(str(path))
    return (
        f"OK — {len(data)} row(s) written to bookmark '{bookmark}' "
        f"in '{path.name}' (python-docx, saved to disk)"
    )


def _docx_find_table_for_bookmark(doc, bookmark_name: str):
    """
    Walk the document XML to find a w:bookmarkStart with the given name,
    then climb the element tree to find its enclosing w:tbl (table).
    Returns a docx Table object or None.
    """
    from docx.oxml.ns import qn
    from docx.table import Table

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body = doc.element.body

    for bk_start in body.iter(qn("w:bookmarkStart")):
        if bk_start.get(qn("w:name")) == bookmark_name:
            # Walk up the tree looking for a table element
            node = bk_start.getparent()
            while node is not None:
                if node.tag == qn("w:tbl"):
                    # Wrap in a docx Table object
                    return Table(node, doc)
                node = node.getparent()
    return None


def _docx_add_row(table):
    """
    Append a new row to a python-docx table by copying the last row's XML
    structure (preserves column widths and basic formatting).
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    last_tr  = table.rows[-1]._tr
    new_tr   = deepcopy(last_tr)
    # Clear text content in the new row
    for tc in new_tr.findall(qn("w:tc")):
        for t in tc.findall(".//" + qn("w:t")):
            t.text = ""
    last_tr.addnext(new_tr)
