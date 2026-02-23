#!/usr/bin/env python3
"""
_word_utils.py — cross-platform Word table population utility for ClipCommand.

Drop this file into your transforms/ folder alongside the word_from_* scripts.
Import it in any transform:
    from _word_utils import update_table, WordUtilsError

Public API
----------
update_table(bookmark, data, heading_rows=1, doc_path=None)
    Write *data* (list of lists) into the first table inside *bookmark*.

    bookmark      : str   — Word bookmark name (no spaces)
    data          : list  — list of rows, each row a list of cell values
    heading_rows  : int   — how many heading rows already exist in the table
                            (these are skipped; data is written below them)
    doc_path      : str   — path to the saved .docx file (required on macOS
                            and Linux; ignored on Windows COM path)

Platform dispatch
-----------------
  Windows  → win32com COM automation  (live active document, no save needed)
  macOS /
  Linux    → python-docx              (file on disk; doc_path required)

macOS note
----------
  JXA and AppleScript automation of Word table cells is unreliable on current
  versions of Word for Mac — cell write operations are silently ignored or
  raise -1728 / -1708 errors regardless of API used. python-docx on a saved
  file is the only reliable approach on non-Windows platforms.

  Workflow on macOS:
    1. Save your document in Word (File → Save).
    2. Close it in Word (or at minimum do not edit it while the transform runs).
    3. Set DOC_PATH in the transform config to the full path of the .docx file.
    4. Run the transform — it writes the data and saves the file.
    5. Reopen the file in Word to see the changes.
"""

import sys
from pathlib import Path


class WordUtilsError(Exception):
    """Raised for all word_utils failures."""


# ─── Public entry point ───────────────────────────────────────────────────────

def update_table(bookmark: str, data: list, heading_rows: int = 1,
                 doc_path: str = None) -> str:
    """
    Write data into the bookmarked Word table.
    Returns a short status string suitable for use as transform() return value.
    """
    if not data:
        raise WordUtilsError("No data rows provided.")

    if sys.platform == "win32":
        return _update_via_com(bookmark, data, heading_rows)
    else:
        if not doc_path:
            raise WordUtilsError(
                "doc_path is required on macOS and Linux.\n"
                "Save your document first, then set DOC_PATH in the transform config.\n"
                "Close the file in Word before running the transform, "
                "then reopen it afterwards to see the changes."
            )
        return _update_via_docx(bookmark, data, heading_rows, doc_path)


# ─── Windows: win32com ────────────────────────────────────────────────────────

def _get_com_word():
    """Return a live Word.Application COM object, clearing the gen_py cache if needed."""
    import shutil
    import os
    import re
    from win32com import client
    try:
        return client.gencache.EnsureDispatch("Word.Application")
    except AttributeError:
        for mod in list(sys.modules):
            if re.match(r"win32com\.gen_py\..+", mod):
                del sys.modules[mod]
        shutil.rmtree(
            os.path.join(os.environ.get("LOCALAPPDATA"), "Temp", "gen_py"),
            ignore_errors=True,
        )
        return client.gencache.EnsureDispatch("Word.Application")


def _update_via_com(bookmark: str, data: list, heading_rows: int) -> str:
    try:
        app = _get_com_word()
        app.Visible       = True
        app.DisplayAlerts = False
        doc               = app.ActiveDocument

        if doc is None:
            raise WordUtilsError("No Word document is currently open.")

        word_range  = doc.Bookmarks(bookmark).Range
        table       = word_range.Tables(1)
        rows_needed = len(data) + heading_rows

        if table.Rows.Count < rows_needed:
            table.Select()
            app.Selection.InsertRowsBelow(NumRows=rows_needed - table.Rows.Count)

        for row_idx, entry in enumerate(data, start=heading_rows + 1):
            for col_idx, value in enumerate(entry, start=1):
                table.Cell(row_idx, col_idx).Range.Text = str(value)

        return (
            f"OK — {len(data)} row(s) written to bookmark '{bookmark}' "
            f"in '{doc.Name}' (Windows COM)"
        )
    except WordUtilsError:
        raise
    except Exception as exc:
        raise WordUtilsError(f"COM error: {exc}") from exc


# ─── macOS / Linux: python-docx ──────────────────────────────────────────────

def _update_via_docx(bookmark: str, data: list, heading_rows: int,
                     doc_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise WordUtilsError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    path = Path(doc_path).expanduser().resolve()
    if not path.exists():
        raise WordUtilsError(f"Document not found: {path}")

    doc   = Document(str(path))
    table = _docx_find_table_for_bookmark(doc, bookmark)

    if table is None:
        raise WordUtilsError(
            f"Bookmark '{bookmark}' not found, or is not inside a table.\n"
            f"Check the bookmark name in Word via Insert → Bookmark."
        )

    # Add rows if the table is too short
    rows_needed = len(data) + heading_rows
    while len(table.rows) < rows_needed:
        _docx_add_row(table)

    # Write cell values
    for row_idx, entry in enumerate(data, start=heading_rows):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(entry):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                # Clear existing content then write
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = str(value)
                else:
                    cell.text = str(value)

    doc.save(str(path))
    return (
        f"OK — {len(data)} row(s) written to bookmark '{bookmark}' "
        f"in '{path.name}' (python-docx, saved to disk)"
    )


def _docx_find_table_for_bookmark(doc, bookmark_name: str):
    """
    Walk document XML to find a w:bookmarkStart with the given name,
    then climb the element tree to its enclosing w:tbl.
    Returns a docx Table object or None.
    """
    from docx.oxml.ns import qn
    from docx.table import Table

    body = doc.element.body
    for bk_start in body.iter(qn("w:bookmarkStart")):
        if bk_start.get(qn("w:name")) == bookmark_name:
            node = bk_start.getparent()
            while node is not None:
                if node.tag == qn("w:tbl"):
                    return Table(node, doc)
                node = node.getparent()
    return None


def _docx_add_row(table):
    """
    Append a new row by deep-copying the last row's XML,
    preserving column widths and basic formatting.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    last_tr = table.rows[-1]._tr
    new_tr  = deepcopy(last_tr)
    for tc in new_tr.findall(qn("w:tc")):
        for t in tc.findall(".//" + qn("w:t")):
            t.text = ""
    last_tr.addnext(new_tr)